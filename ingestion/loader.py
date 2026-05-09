"""Document loaders for PDF, DOCX, TXT, and Markdown files."""

import logging
import os
from pathlib import Path

from core.config import settings
from ingestion.models import Document, normalize_source

logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS = {
    ".pdf": "pdf",
    ".docx": "docx",
    ".txt": "txt",
    ".md": "txt",
}


class DocumentLoader:
    """Load supported files into VaultMind's internal Document model."""

    def __init__(self, raw_data_path: str | None = None) -> None:
        self.raw_data_path = Path(raw_data_path or settings.raw_data_path)
        self.raw_data_path.mkdir(parents=True, exist_ok=True)

    def load(self, file_path: str | Path) -> Document:
        """Load a single document from a file path."""
        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(f"File not found: {path}")

        extension = path.suffix.lower()
        if extension not in SUPPORTED_EXTENSIONS:
            supported = ", ".join(sorted(SUPPORTED_EXTENSIONS))
            raise ValueError(f"Unsupported file type '{extension}'. Supported types: {supported}")

        validate_file_size(str(path))

        file_type = SUPPORTED_EXTENSIONS[extension]
        logger.info("Loading %s file: %s", file_type.upper(), path.name)

        if file_type == "pdf":
            text, metadata = _load_pdf(path)
        elif file_type == "docx":
            text, metadata = _load_docx(path)
        else:
            text, metadata = _load_text(path)

        document = Document(
            text=text.strip(),
            source=normalize_source(path),
            file_type=file_type,
            metadata={
                **metadata,
                "file_name": path.name,
                "file_path": str(path),
                "file_type": file_type,
            },
        )

        logger.info("Loaded '%s' with %s characters", path.name, len(document.text))
        return document

    def load_directory(self, directory: str | Path | None = None) -> list[Document]:
        """Load all supported documents from a directory."""
        dir_path = Path(directory or self.raw_data_path)

        if not dir_path.exists():
            raise FileNotFoundError(f"Directory not found: {dir_path}")

        documents: list[Document] = []
        failed: list[str] = []

        for file_path in sorted(dir_path.iterdir()):
            if not file_path.is_file() or not is_supported_file(file_path.name):
                continue

            try:
                documents.append(self.load(file_path))
            except Exception as exc:
                logger.warning("Failed to load '%s': %s", file_path.name, exc)
                failed.append(file_path.name)

        logger.info(
            "Directory load complete: %s succeeded, %s failed",
            len(documents),
            len(failed),
        )

        if failed:
            logger.warning("Failed files: %s", failed)

        return documents


def load_document(file_path: str | Path) -> Document:
    """Load a supported document using the default loader."""
    return DocumentLoader().load(file_path)


def validate_file_size(file_path: str | Path, max_mb: int | None = None) -> bool:
    """Check that a file is within the allowed upload size."""
    max_mb = settings.max_upload_size_mb if max_mb is None else max_mb
    size_mb = os.path.getsize(file_path) / (1024 * 1024)

    if size_mb > max_mb:
        raise ValueError(f"File too large: {size_mb:.1f}MB. Maximum allowed: {max_mb}MB")

    return True


def is_supported_file(filename: str | Path) -> bool:
    """Return whether a filename has a supported extension."""
    return Path(filename).suffix.lower() in SUPPORTED_EXTENSIONS


def _load_text(path: Path) -> tuple[str, dict[str, object]]:
    text = path.read_text(encoding="utf-8")
    return text, {"page_count": None}


def _load_pdf(path: Path) -> tuple[str, dict[str, object]]:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise ImportError("Install pypdf to load PDF files: pip install pypdf") from exc

    reader = PdfReader(str(path))
    page_texts = [page.extract_text() or "" for page in reader.pages]
    pages = [
        {"page_number": index + 1, "text_length": len(page_text)}
        for index, page_text in enumerate(page_texts)
    ]

    return "\n\n".join(page_texts), {"page_count": len(page_texts), "pages": pages}


def _load_docx(path: Path) -> tuple[str, dict[str, object]]:
    try:
        from docx import Document as DocxDocument
    except ImportError as exc:
        raise ImportError("Install python-docx to load DOCX files: pip install python-docx") from exc

    document = DocxDocument(str(path))
    paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text]

    return "\n".join(paragraphs), {
        "page_count": None,
        "paragraph_count": len(paragraphs),
    }
